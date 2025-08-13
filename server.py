#!/usr/bin/env python3
"""
文档阅读器 MCP 服务器
支持读取 Word (.docx), PDF, TXT, RTF 等格式的文档
"""

import asyncio
import json
import sys
import os
import base64
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse

# MCP imports
from mcp.server import Server
from mcp.server.models import InitializationOptions
import mcp.server.stdio
import mcp.types as types

# 文档处理库
try:
    from docx import Document  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# 图像分析库
try:
    import cv2
    import numpy as np
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False

# 导入简单图表阅读器
try:
    from simple_diagram_reader import SimpleDiagramReader, analyze_single_image
    SIMPLE_DIAGRAM_READER_AVAILABLE = True
except ImportError:
    SIMPLE_DIAGRAM_READER_AVAILABLE = False

# 创建服务器实例
server = Server("document-reader")

@server.list_tools()
async def handle_list_tools() -> List[types.Tool]:
    """列出可用的工具"""
    tools = [
        types.Tool(
            name="read_document",
            description="读取各种格式的文档内容 (Word .docx, PDF, Excel .xlsx/.xls, TXT, RTF)",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件的路径"
                    },
                    "page_range": {
                        "type": "string",
                        "description": "可选：PDF页面范围，格式如 '1-5' 或 '1,3,5'",
                        "default": "all"
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "可选：Excel工作表名称，不指定则读取所有工作表"
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="get_document_info",
            description="获取文档的基本信息（页数、格式、大小等）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件的路径"
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="list_supported_formats",
            description="列出当前支持的文档格式",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        types.Tool(
            name="extract_document_media",
            description="提取文档中的图片和链接信息",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件的路径"
                    },
                    "extract_images": {
                        "type": "boolean",
                        "description": "是否提取图片信息",
                        "default": True
                    },
                    "extract_links": {
                        "type": "boolean",
                        "description": "是否提取链接信息",
                        "default": True
                    },
                    "save_images": {
                        "type": "boolean",
                        "description": "是否保存提取的图片到本地",
                        "default": False
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="read_document_with_media",
            description="读取文档内容并包含媒体元素信息（图片、链接等）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件的路径"
                    },
                    "page_range": {
                        "type": "string",
                        "description": "可选：PDF页面范围，格式如 '1-5' 或 '1,3,5'",
                        "default": "all"
                    },
                    "include_media_info": {
                        "type": "boolean",
                        "description": "是否包含媒体元素信息",
                        "default": True
                    }
                },
                "required": ["file_path"]
            }
        ),

        types.Tool(
            name="read_diagram_content",
            description="直接读取图表内容，无需复杂OCR配置，基于图像结构分析理解图表",
            inputSchema={
                "type": "object",
                "properties": {
                    "image_path": {
                        "type": "string",
                        "description": "图片文件的路径"
                    }
                },
                "required": ["image_path"]
            }
        )
    ]
    return tools

def read_docx_file(file_path: str) -> str:
    """读取Word文档内容"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx 库未安装，无法读取 .docx 文件")

    doc = Document(file_path)
    content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)

    return "\n".join(content)

def read_pdf_file(file_path: str, page_range: str = "all") -> str:
    """读取PDF文档内容"""
    if not PDF_AVAILABLE:
        raise Exception("PyPDF2 库未安装，无法读取 PDF 文件")

    content = []

    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        total_pages = len(pdf_reader.pages)

        # 解析页面范围
        if page_range == "all":
            pages_to_read = range(total_pages)
        else:
            pages_to_read = parse_page_range(page_range, total_pages)

        for page_num in pages_to_read:
            if 0 <= page_num < total_pages:
                page = pdf_reader.pages[page_num]
                content.append(f"=== 第 {page_num + 1} 页 ===\n{page.extract_text()}")

    return "\n\n".join(content)

def parse_page_range(page_range: str, total_pages: int) -> List[int]:
    """解析页面范围字符串"""
    pages = []

    for part in page_range.split(','):
        part = part.strip()
        if '-' in part:
            start, end = map(int, part.split('-'))
            pages.extend(range(start - 1, min(end, total_pages)))
        else:
            page_num = int(part) - 1
            if 0 <= page_num < total_pages:
                pages.append(page_num)

    return sorted(set(pages))

def read_txt_file(file_path: str) -> str:
    """读取文本文件内容"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue

    raise Exception(f"无法使用常见编码读取文件: {file_path}")

def read_rtf_file(file_path: str) -> str:
    """读取RTF文件内容"""
    if not RTF_AVAILABLE:
        raise Exception("striprtf 库未安装，无法读取 RTF 文件")

    with open(file_path, 'r', encoding='utf-8') as file:
        rtf_content = file.read()
        return rtf_to_text(rtf_content)

def read_excel_file(file_path: str, sheet_name: str = None) -> str:
    """读取Excel文件内容"""
    if not EXCEL_AVAILABLE:
        raise Exception("openpyxl 和 pandas 库未安装，无法读取 Excel 文件")

    try:
        # 使用pandas读取Excel文件
        if sheet_name:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content = [f"=== 工作表: {sheet_name} ===\n"]
        else:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            content = []
            
            for sheet in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet)
                content.append(f"=== 工作表: {sheet} ===")
                
                # 转换为字符串格式，保持表格结构
                if not df.empty:
                    # 处理空值
                    df_str = df.fillna('').astype(str)
                    content.append(df_str.to_string(index=False))
                else:
                    content.append("(空工作表)")
                content.append("")  # 添加空行分隔
            
            return "\n".join(content)
        
        # 单个工作表的处理
        if not df.empty:
            df_str = df.fillna('').astype(str)
            content.append(df_str.to_string(index=False))
        else:
            content.append("(空工作表)")
            
        return "\n".join(content)
        
    except Exception as e:
        raise Exception(f"读取Excel文件失败: {str(e)}")

def get_excel_info(file_path: str) -> Dict[str, Any]:
    """获取Excel文件信息"""
    if not EXCEL_AVAILABLE:
        raise Exception("openpyxl 库未安装，无法获取 Excel 文件信息")
    
    try:
        from openpyxl import load_workbook
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        # 加载工作簿
        wb = load_workbook(file_path, read_only=True)
        
        info = {
            "format": "Excel (.xlsx/.xls)",
            "file_size": f"{file_size:,} 字节",
            "sheet_count": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
            "sheets_info": []
        }
        
        # 获取每个工作表的信息
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_info = {
                "name": sheet_name,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "dimensions": f"{ws.max_row} 行 × {ws.max_column} 列"
            }
            info["sheets_info"].append(sheet_info)
        
        wb.close()
        return info
        
    except Exception as e:
        raise Exception(f"获取Excel文件信息失败: {str(e)}")

def extract_excel_media(file_path: str, extract_images: bool = True, extract_links: bool = True, save_images: bool = False) -> Dict[str, Any]:
    """从Excel文件中提取图片和链接信息"""
    if not EXCEL_AVAILABLE:
        raise Exception("openpyxl 库未安装，无法处理 Excel 文件")
    
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as OpenpyxlImage
    import re
    
    result = {"images": [], "links": [], "summary": {}}
    
    try:
        # 加载工作簿
        wb = load_workbook(file_path, data_only=False)
        
        # 提取图片
        if extract_images:
            image_count = 0
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                # 检查工作表中的图片
                if hasattr(ws, '_images') and ws._images:
                    for img in ws._images:
                        try:
                            image_count += 1
                            image_info = {
                                "sheet": sheet_name,
                                "filename": f"excel_image_{image_count}.{img.format.lower()}",
                                "format": img.format,
                                "anchor": str(img.anchor) if hasattr(img, 'anchor') else "未知位置"
                            }
                            
                            # 如果需要保存图片
                            if save_images and PILLOW_AVAILABLE:
                                try:
                                    # 创建保存目录
                                    save_dir = Path("extracted_images")
                                    save_dir.mkdir(exist_ok=True)
                                    
                                    # 保存图片
                                    image_path = save_dir / image_info["filename"]
                                    with open(image_path, 'wb') as f:
                                        f.write(img._data())
                                    
                                    image_info["saved_path"] = str(image_path)
                                    image_info["file_size"] = len(img._data())
                                    
                                    # 获取图片尺寸
                                    try:
                                        from PIL import Image
                                        with Image.open(image_path) as pil_img:
                                            image_info["dimensions"] = pil_img.size
                                    except:
                                        pass
                                        
                                except Exception as e:
                                    image_info["save_error"] = str(e)
                            
                            result["images"].append(image_info)
                            
                        except Exception as e:
                            result["images"].append({
                                "sheet": sheet_name,
                                "error": f"图片处理失败: {str(e)}"
                            })
        
        # 提取链接
        if extract_links:
            url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+')
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                # 检查单元格中的超链接
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.hyperlink:
                            try:
                                link_info = {
                                    "sheet": sheet_name,
                                    "cell": f"{cell.coordinate}",
                                    "url": cell.hyperlink.target,
                                    "display_text": str(cell.value) if cell.value else "",
                                    "type": "hyperlink"
                                }
                                
                                # 验证链接有效性
                                if REQUESTS_AVAILABLE and link_info["url"].startswith(('http://', 'https://')):
                                    try:
                                        import requests
                                        response = requests.head(link_info["url"], timeout=5, allow_redirects=True)
                                        link_info["status_code"] = response.status_code
                                        link_info["accessible"] = response.status_code < 400
                                    except:
                                        link_info["accessible"] = False
                                        link_info["status_code"] = "连接失败"
                                
                                result["links"].append(link_info)
                                
                            except Exception as e:
                                result["links"].append({
                                    "sheet": sheet_name,
                                    "cell": f"{cell.coordinate}",
                                    "error": f"链接处理失败: {str(e)}"
                                })
                        
                        # 检查单元格文本中的URL
                        elif cell.value and isinstance(cell.value, str):
                            urls = url_pattern.findall(str(cell.value))
                            for url in urls:
                                try:
                                    link_info = {
                                        "sheet": sheet_name,
                                        "cell": f"{cell.coordinate}",
                                        "url": url,
                                        "display_text": str(cell.value),
                                        "type": "text_url"
                                    }
                                    
                                    # 验证链接有效性
                                    if REQUESTS_AVAILABLE:
                                        try:
                                            import requests
                                            response = requests.head(url, timeout=5, allow_redirects=True)
                                            link_info["status_code"] = response.status_code
                                            link_info["accessible"] = response.status_code < 400
                                        except:
                                            link_info["accessible"] = False
                                            link_info["status_code"] = "连接失败"
                                    
                                    result["links"].append(link_info)
                                    
                                except Exception as e:
                                    result["links"].append({
                                        "sheet": sheet_name,
                                        "cell": f"{cell.coordinate}",
                                        "url": url,
                                        "error": f"链接处理失败: {str(e)}"
                                    })
        
        # 生成摘要
        result["summary"] = {
            "image_count": len([img for img in result["images"] if "error" not in img]),
            "link_count": len([link for link in result["links"] if "error" not in link]),
            "image_errors": len([img for img in result["images"] if "error" in img]),
            "link_errors": len([link for link in result["links"] if "error" in link])
        }
        
        wb.close()
        return result
        
    except Exception as e:
        raise Exception(f"Excel媒体提取失败: {str(e)}")

def read_excel_with_media(file_path: str, sheet_name: str = None) -> Tuple[str, Dict[str, Any]]:
    """读取Excel文档内容并提取媒体信息"""
    if not EXCEL_AVAILABLE:
        raise Exception("openpyxl 和 pandas 库未安装，无法读取 Excel 文件")
    
    # 读取文档内容
    content = read_excel_file(file_path, sheet_name)
    
    # 提取媒体信息
    try:
        media_data = extract_excel_media(file_path, extract_images=True, extract_links=True, save_images=False)
        media_info = {
            "images": media_data.get("images", []),
            "links": media_data.get("links", []),
            "summary": media_data.get("summary", {})
        }
        return content, media_info
    except Exception as e:
        # 如果媒体提取失败，返回空的媒体信息
        return content, {"images": [], "links": [], "summary": {"error": str(e)}}

def analyze_flowchart_image_from_bytes(image_bytes: bytes) -> dict:
    """分析流程图图片，基于OpenCV的基础结构分析"""
    result = {"text": "", "nodes": 0, "edges": 0}
    try:
        if not OPENCV_AVAILABLE:
            result["error"] = "OpenCV不可用"
            return result
            
        # 读取图片
        import io
        file_bytes = np.asarray(bytearray(image_bytes), dtype=np.uint8)
        img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if img is None:
            result["error"] = "无法读取图片"
            return result
            
        # 灰度处理
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # 边缘检测
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        
        # 检测圆形节点
        circles = cv2.HoughCircles(gray, cv2.HOUGH_GRADIENT, dp=1, minDist=40, param1=50, param2=30, minRadius=10, maxRadius=80)
        if circles is not None:
            result["nodes"] = len(circles[0])
            
        # 检测直线（连线）
        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=80, minLineLength=40, maxLineGap=10)
        if lines is not None:
            result["edges"] = len(lines)
            
    except Exception as e:
        result["error"] = f"流程图解析失败: {str(e)}"
    return result

def format_simple_diagram_result(result: Dict[str, Any]) -> str:
    """格式化简单图表分析结果"""
    if "error" in result:
        return f"分析失败: {result['error']}"
    
    result_text = "=== 图表内容分析 ===\n\n"
    
    # 文件信息
    if "file_info" in result:
        info = result["file_info"]
        result_text += f"📊 文件信息:\n"
        result_text += f"- 文件名: {info.get('filename', '未知')}\n"
        result_text += f"- 尺寸: {info.get('dimensions', '未知')}\n"
        result_text += f"- 大小: {info.get('size', '未知')}\n\n"
    
    # 图表解释
    if "interpretation" in result:
        interp = result["interpretation"]
        result_text += f"🎯 图表类型: {interp.get('predicted_type', '未知')}\n"
        result_text += f"📈 置信度: {interp.get('confidence', 0):.1%}\n\n"
        result_text += f"📝 内容描述:\n{interp.get('content_description', '无描述')}\n\n"
        
        tech_elements = interp.get('technical_elements', [])
        if tech_elements:
            result_text += f"🔧 技术元素: {', '.join(tech_elements)}\n\n"
    
    # 结构分析
    if "analysis" in result:
        analysis = result["analysis"]
        shapes = analysis.get("shapes", {})
        result_text += f"🏗️ 结构分析:\n"
        result_text += f"- 矩形框: {shapes.get('rectangles', 0)} 个\n"
        result_text += f"- 圆形: {shapes.get('circles', 0)} 个\n"
        result_text += f"- 连接线: {shapes.get('lines', 0)} 条\n"
        result_text += f"- 复杂度评分: {analysis.get('complexity', 0)}\n"
        result_text += f"- 主要方向: {analysis.get('dominant_direction', '未知')}\n\n"
        
        layout = analysis.get("layout", {})
        if layout:
            result_text += f"📐 布局特征:\n"
            result_text += f"- 宽高比: {layout.get('aspect_ratio', 0):.2f}\n"
            result_text += f"- 方向: {layout.get('primary_orientation', '未知')}\n"
    
    return result_text



def extract_docx_media(file_path: str, extract_images: bool = True, extract_links: bool = True, save_images: bool = False) -> Dict[str, Any]:
    """从Word文档中提取图片和链接信息"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx 库未安装，无法处理 .docx 文件")

    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    doc = Document(file_path)
    result = {"images": [], "links": [], "summary": {}}

    # 提取图片
    if extract_images and PILLOW_AVAILABLE:
        try:
            # 从文档关系中获取图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_data = image_part.blob

                    # 获取图片信息
                    try:
                        from io import BytesIO
                        image = Image.open(BytesIO(image_data))

                        image_info = {
                            "filename": rel.target_ref.split('/')[-1],
                            "format": image.format,
                            "size": image.size,
                            "mode": image.mode,
                            "data_size": len(image_data)
                        }

                        # 如果需要保存图片
                        if save_images:
                            output_dir = Path(file_path).parent / "extracted_images"
                            output_dir.mkdir(exist_ok=True)
                            image_path = output_dir / image_info["filename"]
                            with open(image_path, 'wb') as f:
                                f.write(image_data)
                            image_info["saved_path"] = str(image_path)

                        # 新增流程图内容解析
                        try:
                            flowchart_info = analyze_flowchart_image_from_bytes(image_data)
                            image_info["flowchart_analysis"] = flowchart_info
                        except Exception as e:
                            image_info["flowchart_analysis_error"] = str(e)

                        result["images"].append(image_info)
                    except Exception as e:
                        result["images"].append({
                            "filename": rel.target_ref.split('/')[-1],
                            "error": f"无法处理图片: {str(e)}"
                        })
        except Exception as e:
            result["images"] = [{"error": f"提取图片时出错: {str(e)}"}]

    # 提取链接
    if extract_links:
        try:
            # 从段落中提取超链接
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.element.tag.endswith('hyperlink') or any('hyperlink' in str(child.tag) for child in run.element):
                        # 这是一个简化的链接检测，实际实现可能需要更复杂的XML解析
                        pass

            # 使用正则表达式从文本中提取URL
            full_text = "\n".join([p.text for p in doc.paragraphs])
            url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
            urls = re.findall(url_pattern, full_text)

            for url in set(urls):  # 去重
                link_info = {
                    "url": url,
                    "domain": urlparse(url).netloc,
                    "scheme": urlparse(url).scheme
                }

                # 如果有requests库，检查链接有效性
                if REQUESTS_AVAILABLE:
                    try:
                        response = requests.head(url, timeout=5, allow_redirects=True)
                        link_info["status_code"] = response.status_code
                        link_info["accessible"] = response.status_code < 400
                    except:
                        link_info["accessible"] = False
                        link_info["status_code"] = None

                result["links"].append(link_info)

        except Exception as e:
            result["links"] = [{"error": f"提取链接时出错: {str(e)}"}]

    # 生成摘要
    result["summary"] = {
        "total_images": len([img for img in result["images"] if "error" not in img]),
        "total_links": len([link for link in result["links"] if "error" not in link]),
        "images_with_errors": len([img for img in result["images"] if "error" in img]),
        "links_with_errors": len([link for link in result["links"] if "error" in link])
    }

    return result

def extract_pdf_links(file_path: str) -> List[Dict[str, Any]]:
    """从PDF文档中提取链接信息"""
    if not PDF_AVAILABLE:
        raise Exception("PyPDF2 库未安装，无法处理 PDF 文件")

    links = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                # 提取页面文本中的URL
                text = page.extract_text()
                url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
                urls = re.findall(url_pattern, text)

                for url in urls:
                    link_info = {
                        "url": url,
                        "page": page_num + 1,
                        "domain": urlparse(url).netloc,
                        "scheme": urlparse(url).scheme
                    }

                    # 检查链接有效性
                    if REQUESTS_AVAILABLE:
                        try:
                            response = requests.head(url, timeout=5, allow_redirects=True)
                            link_info["status_code"] = response.status_code
                            link_info["accessible"] = response.status_code < 400
                        except:
                            link_info["accessible"] = False
                            link_info["status_code"] = None

                    links.append(link_info)

    except Exception as e:
        return [{"error": f"提取PDF链接时出错: {str(e)}"}]

    return links

def read_docx_with_media(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """读取Word文档内容并提取媒体信息"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx 库未安装，无法读取 .docx 文件")

    doc = Document(file_path)
    content = []
    media_info = {"images": [], "links": []}

    # 提取文本内容
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)

    # 提取媒体信息
    try:
        media_data = extract_docx_media(file_path, extract_images=True, extract_links=True, save_images=False)
        media_info = {
            "images": media_data.get("images", []),
            "links": media_data.get("links", []),
            "summary": media_data.get("summary", {})
        }
    except Exception as e:
        media_info["error"] = f"提取媒体信息时出错: {str(e)}"

    return "\n".join(content), media_info

def get_file_info(file_path: str) -> Dict[str, Any]:
    """获取文件基本信息"""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    info = {
        "文件名": path.name,
        "文件大小": f"{path.stat().st_size / 1024:.2f} KB",
        "文件格式": path.suffix.lower(),
        "绝对路径": str(path.absolute())
    }

    # 根据文件类型添加特定信息
    if path.suffix.lower() == '.docx' and DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            info["段落数"] = len(doc.paragraphs)
        except:
            pass

    elif path.suffix.lower() == '.pdf' and PDF_AVAILABLE:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                info["页数"] = len(pdf_reader.pages)
        except:
            pass
    
    elif path.suffix.lower() in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
        try:
            excel_info = get_excel_info(file_path)
            info.update(excel_info)
        except:
            pass

    return info

@server.call_tool()
async def handle_call_tool(
    name: str, arguments: Dict[str, Any]
) -> List[types.TextContent]:
    """处理工具调用"""

    if name == "read_document":
        file_path = arguments.get("file_path")
        page_range = arguments.get("page_range", "all")

        if not file_path:
            return [types.TextContent(
                type="text",
                text="错误：必须提供文件路径"
            )]

        try:
            path = Path(file_path)

            if not path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"错误：文件不存在 - {file_path}"
                )]

            file_ext = path.suffix.lower()

            if file_ext == '.docx':
                content = read_docx_file(file_path)
            elif file_ext == '.pdf':
                content = read_pdf_file(file_path, page_range)
            elif file_ext in ['.xlsx', '.xls']:
                sheet_name = arguments.get("sheet_name")
                content = read_excel_file(file_path, sheet_name)
            elif file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css']:
                content = read_txt_file(file_path)
            elif file_ext == '.rtf':
                content = read_rtf_file(file_path)
            else:
                return [types.TextContent(
                    type="text",
                    text=f"错误：不支持的文件格式 - {file_ext}"
                )]

            return [types.TextContent(
                type="text",
                text=f"文档内容 ({path.name}):\n\n{content}"
            )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"读取文档时出错: {str(e)}"
            )]

    elif name == "get_document_info":
        file_path = arguments.get("file_path")

        if not file_path:
            return [types.TextContent(
                type="text",
                text="错误：必须提供文件路径"
            )]

        try:
            info = get_file_info(file_path)
            info_text = "\n".join([f"{k}: {v}" for k, v in info.items()])

            return [types.TextContent(
                type="text",
                text=f"文档信息:\n{info_text}"
            )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"获取文档信息时出错: {str(e)}"
            )]

    elif name == "list_supported_formats":
        formats = {
            ".docx": "Word文档" + (" ✓" if DOCX_AVAILABLE else " ✗ (需要 python-docx)"),
            ".pdf": "PDF文档" + (" ✓" if PDF_AVAILABLE else " ✗ (需要 PyPDF2)"),
            ".xlsx/.xls": "Excel文档" + (" ✓" if EXCEL_AVAILABLE else " ✗ (需要 openpyxl, pandas)"),
            ".txt": "纯文本文件 ✓",
            ".md": "Markdown文件 ✓",
            ".rtf": "RTF文档" + (" ✓" if RTF_AVAILABLE else " ✗ (需要 striprtf)"),
            ".py/.js/.html/.css": "代码文件 ✓"
        }

        # 添加媒体处理能力信息
        media_support = {
            "图片处理": " ✓" if PILLOW_AVAILABLE else " ✗ (需要 Pillow)",
            "链接验证": " ✓" if REQUESTS_AVAILABLE else " ✗ (需要 requests)"
        }

        format_text = "\n".join([f"{ext}: {desc}" for ext, desc in formats.items()])
        media_text = "\n".join([f"{feature}: {status}" for feature, status in media_support.items()])

        return [types.TextContent(
            type="text",
            text=f"支持的文档格式:\n{format_text}\n\n媒体处理能力:\n{media_text}"
        )]

    elif name == "extract_document_media":
        file_path = arguments.get("file_path")
        extract_images = arguments.get("extract_images", True)
        extract_links = arguments.get("extract_links", True)
        save_images = arguments.get("save_images", False)

        if not file_path:
            return [types.TextContent(
                type="text",
                text="错误：必须提供文件路径"
            )]

        try:
            path = Path(file_path)

            if not path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"错误：文件不存在 - {file_path}"
                )]

            file_ext = path.suffix.lower()

            if file_ext == '.docx':
                media_data = extract_docx_media(file_path, extract_images, extract_links, save_images)

                result_text = f"文档媒体信息 ({path.name}):\n\n"

                # 摘要信息
                summary = media_data.get("summary", {})
                result_text += f"摘要:\n"
                result_text += f"- 图片总数: {summary.get('total_images', 0)}\n"
                result_text += f"- 链接总数: {summary.get('total_links', 0)}\n"
                result_text += f"- 图片处理错误: {summary.get('images_with_errors', 0)}\n"
                result_text += f"- 链接处理错误: {summary.get('links_with_errors', 0)}\n\n"

                # 图片信息
                if extract_images and media_data.get("images"):
                    result_text += "图片信息:\n"
                    for i, img in enumerate(media_data["images"], 1):
                        if "error" in img:
                            result_text += f"{i}. 错误: {img['error']}\n"
                        else:
                            result_text += f"{i}. {img.get('filename', '未知文件名')}\n"
                            result_text += f"   - 格式: {img.get('format', '未知')}\n"
                            result_text += f"   - 尺寸: {img.get('size', '未知')}\n"
                            result_text += f"   - 大小: {img.get('data_size', 0)} 字节\n"
                            if "saved_path" in img:
                                result_text += f"   - 保存路径: {img['saved_path']}\n"
                            if "flowchart_analysis" in img:
                                flowchart = img["flowchart_analysis"]
                                result_text += f"   - 流程图分析:\n"
                                result_text += f"     - 文本: {flowchart.get('text', '无')}\n"
                                result_text += f"     - 节点数: {flowchart.get('nodes', 0)}\n"
                                result_text += f"     - 连线数: {flowchart.get('edges', 0)}\n"
                            if "flowchart_analysis_error" in img:
                                result_text += f"   - 流程图分析错误: {img['flowchart_analysis_error']}\n"
                    result_text += "\n"

                # 链接信息
                if extract_links and media_data.get("links"):
                    result_text += "链接信息:\n"
                    for i, link in enumerate(media_data["links"], 1):
                        if "error" in link:
                            result_text += f"{i}. 错误: {link['error']}\n"
                        else:
                            result_text += f"{i}. {link['url']}\n"
                            result_text += f"   - 域名: {link.get('domain', '未知')}\n"
                            result_text += f"   - 协议: {link.get('scheme', '未知')}\n"
                            if "accessible" in link:
                                status = "可访问" if link["accessible"] else "不可访问"
                                result_text += f"   - 状态: {status}"
                                if link.get("status_code"):
                                    result_text += f" (HTTP {link['status_code']})"
                                result_text += "\n"

                return [types.TextContent(
                    type="text",
                    text=result_text
                )]

            elif file_ext in ['.xlsx', '.xls']:
                media_data = extract_excel_media(file_path, extract_images, extract_links, save_images)

                result_text = f"Excel媒体信息 ({path.name}):\n\n"

                # 摘要信息
                summary = media_data.get("summary", {})
                result_text += f"摘要:\n"
                result_text += f"- 图片总数: {summary.get('image_count', 0)}\n"
                result_text += f"- 链接总数: {summary.get('link_count', 0)}\n"
                result_text += f"- 图片处理错误: {summary.get('image_errors', 0)}\n"
                result_text += f"- 链接处理错误: {summary.get('link_errors', 0)}\n\n"

                # 图片信息
                if extract_images and media_data.get("images"):
                    result_text += "图片信息:\n"
                    for i, img in enumerate(media_data["images"], 1):
                        if "error" in img:
                            result_text += f"{i}. 错误: {img['error']}\n"
                        else:
                            result_text += f"{i}. {img.get('filename', '未知文件名')}\n"
                            result_text += f"   - 工作表: {img.get('sheet', '未知')}\n"
                            result_text += f"   - 格式: {img.get('format', '未知')}\n"
                            result_text += f"   - 位置: {img.get('anchor', '未知')}\n"
                            if "dimensions" in img:
                                result_text += f"   - 尺寸: {img['dimensions']}\n"
                            if "file_size" in img:
                                result_text += f"   - 大小: {img['file_size']} 字节\n"
                            if "saved_path" in img:
                                result_text += f"   - 保存路径: {img['saved_path']}\n"
                    result_text += "\n"

                # 链接信息
                if extract_links and media_data.get("links"):
                    result_text += "链接信息:\n"
                    for i, link in enumerate(media_data["links"], 1):
                        if "error" in link:
                            result_text += f"{i}. 错误: {link['error']}\n"
                        else:
                            result_text += f"{i}. {link['url']}\n"
                            result_text += f"   - 工作表: {link.get('sheet', '未知')}\n"
                            result_text += f"   - 单元格: {link.get('cell', '未知')}\n"
                            result_text += f"   - 类型: {link.get('type', '未知')}\n"
                            if link.get('display_text'):
                                result_text += f"   - 显示文本: {link['display_text'][:50]}{'...' if len(link['display_text']) > 50 else ''}\n"
                            if "accessible" in link:
                                status = "可访问" if link["accessible"] else "不可访问"
                                result_text += f"   - 状态: {status}"
                                if link.get("status_code"):
                                    result_text += f" (HTTP {link['status_code']})"
                                result_text += "\n"

                return [types.TextContent(
                    type="text",
                    text=result_text
                )]

            elif file_ext == '.pdf':
                if extract_links:
                    links = extract_pdf_links(file_path)
                    result_text = f"PDF链接信息 ({path.name}):\n\n"

                    if links and not any("error" in link for link in links):
                        result_text += f"找到 {len(links)} 个链接:\n\n"
                        for i, link in enumerate(links, 1):
                            result_text += f"{i}. {link['url']} (第{link['page']}页)\n"
                            result_text += f"   - 域名: {link.get('domain', '未知')}\n"
                            if "accessible" in link:
                                status = "可访问" if link["accessible"] else "不可访问"
                                result_text += f"   - 状态: {status}"
                                if link.get("status_code"):
                                    result_text += f" (HTTP {link['status_code']})"
                                result_text += "\n"
                    else:
                        result_text += "未找到链接或处理时出错\n"
                        for link in links:
                            if "error" in link:
                                result_text += f"错误: {link['error']}\n"

                    return [types.TextContent(
                        type="text",
                        text=result_text
                    )]
                else:
                    return [types.TextContent(
                        type="text",
                        text="PDF文档目前只支持链接提取，请启用 extract_links 参数"
                    )]
            else:
                return [types.TextContent(
                    type="text",
                    text=f"错误：文件格式 {file_ext} 暂不支持媒体提取功能"
                )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"提取媒体信息时出错: {str(e)}"
            )]



    elif name == "read_diagram_content":
        image_path = arguments.get("image_path")

        if not image_path:
            return [types.TextContent(
                type="text",
                text="错误：必须提供图片路径"
            )]

        try:
            path = Path(image_path)

            if not path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"错误：图片文件不存在 - {image_path}"
                )]

            # 检查是否为图片文件
            image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.emf', '.wmf']
            if path.suffix.lower() not in image_extensions:
                return [types.TextContent(
                    type="text",
                    text=f"错误：不支持的图片格式 - {path.suffix}"
                )]

            # 使用简单图表阅读器
            if SIMPLE_DIAGRAM_READER_AVAILABLE:
                result = analyze_single_image(image_path)
                
                if "error" in result:
                    return [types.TextContent(
                        type="text",
                        text=f"图表分析失败: {result['error']}"
                    )]
                
                # 格式化结果
                formatted_result = format_simple_diagram_result(result)
                
                return [types.TextContent(
                    type="text",
                    text=formatted_result
                )]
            else:
                return [types.TextContent(
                    type="text",
                    text="错误：简单图表阅读器不可用"
                )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"读取图表内容时出错: {str(e)}"
            )]

    elif name == "read_document_with_media":
        file_path = arguments.get("file_path")
        page_range = arguments.get("page_range", "all")
        include_media_info = arguments.get("include_media_info", True)

        if not file_path:
            return [types.TextContent(
                type="text",
                text="错误：必须提供文件路径"
            )]

        try:
            path = Path(file_path)

            if not path.exists():
                return [types.TextContent(
                    type="text",
                    text=f"错误：文件不存在 - {file_path}"
                )]

            file_ext = path.suffix.lower()

            if file_ext == '.docx':
                if include_media_info:
                    content, media_info = read_docx_with_media(file_path)

                    result_text = f"文档内容 ({path.name}):\n\n{content}\n\n"

                    # 添加媒体信息
                    if "error" not in media_info:
                        summary = media_info.get("summary", {})
                        result_text += "=== 媒体信息 ===\n"
                        result_text += f"图片数量: {summary.get('total_images', 0)}\n"
                        result_text += f"链接数量: {summary.get('total_links', 0)}\n"

                        if media_info.get("images"):
                            result_text += "\n图片列表:\n"
                            for i, img in enumerate(media_info["images"], 1):
                                if "error" not in img:
                                    result_text += f"{i}. {img.get('filename', '未知')} ({img.get('format', '未知')}, {img.get('size', '未知')})\n"

                        if media_info.get("links"):
                            result_text += "\n链接列表:\n"
                            for i, link in enumerate(media_info["links"], 1):
                                if "error" not in link:
                                    result_text += f"{i}. {link['url']}\n"
                    else:
                        result_text += f"媒体信息提取错误: {media_info['error']}\n"

                    return [types.TextContent(
                        type="text",
                        text=result_text
                    )]
                else:
                    content = read_docx_file(file_path)
                    return [types.TextContent(
                        type="text",
                        text=f"文档内容 ({path.name}):\n\n{content}"
                    )]

            elif file_ext == '.pdf':
                content = read_pdf_file(file_path, page_range)

                if include_media_info:
                    # 为PDF添加链接信息
                    try:
                        links = extract_pdf_links(file_path)
                        result_text = f"文档内容 ({path.name}):\n\n{content}\n\n"

                        if links and not any("error" in link for link in links):
                            result_text += "=== 链接信息 ===\n"
                            result_text += f"找到 {len(links)} 个链接:\n"
                            for i, link in enumerate(links, 1):
                                result_text += f"{i}. {link['url']} (第{link['page']}页)\n"

                        return [types.TextContent(
                            type="text",
                            text=result_text
                        )]
                    except:
                        # 如果链接提取失败，只返回内容
                        return [types.TextContent(
                            type="text",
                            text=f"文档内容 ({path.name}):\n\n{content}"
                        )]
                else:
                    return [types.TextContent(
                        type="text",
                        text=f"文档内容 ({path.name}):\n\n{content}"
                    )]

            elif file_ext in ['.xlsx', '.xls']:
                sheet_name = arguments.get("sheet_name")
                
                if include_media_info:
                    content, media_info = read_excel_with_media(file_path, sheet_name)

                    result_text = f"文档内容 ({path.name}):\n\n{content}\n\n"

                    # 添加媒体信息
                    if "error" not in media_info.get("summary", {}):
                        summary = media_info.get("summary", {})
                        result_text += "=== 媒体信息 ===\n"
                        result_text += f"图片数量: {summary.get('image_count', 0)}\n"
                        result_text += f"链接数量: {summary.get('link_count', 0)}\n"

                        if media_info.get("images"):
                            result_text += "\n图片列表:\n"
                            for i, img in enumerate(media_info["images"], 1):
                                if "error" not in img:
                                    result_text += f"{i}. {img.get('filename', '未知')} ({img.get('format', '未知')}, 工作表: {img.get('sheet', '未知')})\n"

                        if media_info.get("links"):
                            result_text += "\n链接列表:\n"
                            for i, link in enumerate(media_info["links"], 1):
                                if "error" not in link:
                                    result_text += f"{i}. {link['url']} (工作表: {link.get('sheet', '未知')}, 单元格: {link.get('cell', '未知')})\n"
                    else:
                        result_text += f"媒体信息提取错误: {media_info.get('summary', {}).get('error', '未知错误')}\n"

                    return [types.TextContent(
                        type="text",
                        text=result_text
                    )]
                else:
                    content = read_excel_file(file_path, sheet_name)
                    return [types.TextContent(
                        type="text",
                        text=f"文档内容 ({path.name}):\n\n{content}"
                    )]

            elif file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css']:
                content = read_txt_file(file_path)

                if include_media_info:
                    # 从文本文件中提取链接
                    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
                    urls = re.findall(url_pattern, content)

                    result_text = f"文档内容 ({path.name}):\n\n{content}\n\n"

                    if urls:
                        result_text += "=== 链接信息 ===\n"
                        result_text += f"找到 {len(set(urls))} 个链接:\n"
                        for i, url in enumerate(set(urls), 1):
                            result_text += f"{i}. {url}\n"

                    return [types.TextContent(
                        type="text",
                        text=result_text
                    )]
                else:
                    return [types.TextContent(
                        type="text",
                        text=f"文档内容 ({path.name}):\n\n{content}"
                    )]

            elif file_ext == '.rtf':
                content = read_rtf_file(file_path)
                return [types.TextContent(
                    type="text",
                    text=f"文档内容 ({path.name}):\n\n{content}"
                )]
            else:
                return [types.TextContent(
                    type="text",
                    text=f"错误：不支持的文件格式 - {file_ext}"
                )]

        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"读取文档时出错: {str(e)}"
            )]

    else:
        return [types.TextContent(
            type="text",
            text=f"未知工具: {name}"
        )]

# 尝试导入NotificationOptions，如果失败则使用替代方案
try:
    from mcp.server.lowlevel.server import NotificationOptions
    notification_options = NotificationOptions()
except ImportError:
    # 如果导入失败，使用None或空字典
    notification_options = None

async def main():
    """使用标准输入输出运行服务器"""
    try:
        async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
            # 构建capabilities参数
            capabilities_kwargs = {"experimental_capabilities": {}}
            if notification_options is not None:
                capabilities_kwargs["notification_options"] = notification_options

            await server.run(
                read_stream,
                write_stream,
                InitializationOptions(
                    server_name="document-reader",
                    server_version="1.1.0",
                    capabilities=server.get_capabilities(**capabilities_kwargs),
                ),
            )
    except Exception as e:
        print(f"服务器启动失败: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    asyncio.run(main())